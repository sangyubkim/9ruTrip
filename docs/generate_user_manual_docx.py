#!/usr/bin/env python3
"""Generate 9ruTrip Korean end-user manual (.md + .docx) from current product UX.

Usage:
  C:\\Users\\Intellian\\AppData\\Local\\Programs\\Python\\Python312\\python.exe docs\\generate_user_manual_docx.py

Outputs:
  docs/9ruTrip-사용자매뉴얼.md
  docs/9ruTrip-사용자매뉴얼.docx

Brand: #0c4a6e · Font: Malgun Gothic
Based on shipped features through UX B7–B11, JP transit deeplinks, etc.
Do not invent unshipped features. Never embed real API keys.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOCS = Path(__file__).resolve().parent
OUT_MD = DOCS / "9ruTrip-사용자매뉴얼.md"
OUT_DOCX = DOCS / "9ruTrip-사용자매뉴얼.docx"

C_PRIMARY = RGBColor(0x0C, 0x4A, 0x6E)  # #0c4a6e
C_INK = RGBColor(0x0F, 0x17, 0x2A)
C_MUTED = RGBColor(0x64, 0x74, 0x8B)
FONT = "Malgun Gothic"
REF_COMMIT = "81d7276"  # UX B7–B11
TODAY = date.today().isoformat()

# ---------------------------------------------------------------------------
# Shared content model (keeps .md and .docx in sync)
# ---------------------------------------------------------------------------
# Block kinds: h1, h2, h3, p, ul, ol, note, code, table


ContentBlock = dict  # typed loosely for clarity


def build_blocks() -> list[ContentBlock]:
    return [
        {
            "kind": "h1",
            "text": "9ruTrip 사용자 매뉴얼",
        },
        {
            "kind": "p",
            "text": (
                f"모바일 AI 여행 플래너 · Plan → Trip → After\n"
                f"대상: 여행자 · 현장 운영(실기기/에뮬) · 발행 워크플로 사용자\n"
                f"기준 구현: UX B7–B11, JP transit 딥링크 등 (커밋 `{REF_COMMIT}` 포함)\n"
                f"문서 갱신: {TODAY}"
            ),
        },
        {
            "kind": "note",
            "text": (
                "이 문서는 현재 저장된(shipped) 기능만 설명합니다. "
                "국내 네이버 맵 풀 SDK, 네이티브 SMS 인박스 자동 읽기, 계정 클라우드 동기화 등은 "
                "아직 보류 상태이며 본 매뉴얼에 포함하지 않습니다."
            ),
        },
        # ---- 1 ----
        {"kind": "h2", "text": "1. 소개 — 9ruTrip이란"},
        {
            "kind": "p",
            "text": (
                "9ruTrip은 Gemini로 초기 일정을 만들고, 현장에서 동선·경비·기록을 관리한 뒤, "
                "여행 후 BlogDraft / WordPress 발행까지 이어 주는 **Expo Android-first** 여행 앱입니다. "
                "해외 MVP 도시는 **도쿄(기본)** 과 **오사카(선택)** 이며, 지도는 Google Maps를 사용합니다."
            ),
        },
        {"kind": "h3", "text": "Plan / Trip / After"},
        {
            "kind": "ul",
            "items": [
                "**Plan(계획중)** — 도시·박/일/인원으로 AI 일정 생성. Day 타임라인, ≡ DnD, 필터, 지도, 이동 비교, 숙소, 동선 최적화, 체크리스트·날씨.",
                "**Trip(여행중)** — 현장 모드·다음 액션(길안내·완료·재루트), 가이드 알람, GPS 이탈 배너, 사진·리뷰·경비.",
                "**After(완료)** — 계획 vs 실제 요약·인사이트, BlogDraft 내보내기, WordPress 임시글/게시(자격증명 필요).",
            ],
        },
        {
            "kind": "p",
            "text": (
                "여행 상태는 홈·Plan 상단에 **계획중 / 여행중 / 완료** 칩으로 표시됩니다. "
                "데이터는 기기 **AsyncStorage**에 로컬 저장됩니다(계정 동기화 없음)."
            ),
        },
        # ---- 2 ----
        {"kind": "h2", "text": "2. 준비 — API · 환경 변수 · 앱 실행"},
        {"kind": "h3", "text": "2.1 저장소 설치"},
        {
            "kind": "ol",
            "items": [
                "`cd D:\\01_Project\\9ruTrip` 후 `npm install`",
                "`npm run shared:build` 로 `@9rutrip/shared` 빌드",
                "API: `apps/api/.env` 준비 후 `npm run api` (기본 포트 **3011**)",
                "모바일: `cd apps/mobile` → `npm install` → `npm start` → Android(`a` 또는 `npm run android`)",
            ],
        },
        {"kind": "h3", "text": "2.2 API 환경 변수 (`apps/api/.env`)"},
        {
            "kind": "p",
            "text": (
                "`.env.example`을 복사해 채웁니다. **실제 비밀 키는 커밋하지 마세요.** "
                "아래는 형식 예시입니다."
            ),
        },
        {
            "kind": "code",
            "text": (
                "PORT=3011\n"
                "GEMINI_API_KEY=your_gemini_key_here\n"
                "GEMINI_MODEL=gemini-flash-lite-latest\n"
                "\n"
                "# WordPress (선택) — Application Password\n"
                "WP_SITE_URL=https://example.com\n"
                "WP_USERNAME=your_wp_user\n"
                "WP_APP_PASSWORD=xxxx xxxx xxxx xxxx\n"
                "\n"
                "# Directions 서버 키 (선택). 없으면 하버사인 폴백\n"
                "# GOOGLE_MAPS_API_KEY=your_server_maps_key\n"
                "\n"
                "# 일본 대중교통 파트너 (선택). 미설정 시 추정 + Yahoo/Google 딥링크\n"
                "# JP_TRANSIT_PROVIDER=navitime\n"
                "# NAVITIME_API_KEY=\n"
                "# NAVITIME_API_HOST=\n"
            ),
        },
        {
            "kind": "note",
            "text": (
                "Gemini가 없어도 도쿄 폴백 코스로 일정 생성은 가능합니다. "
                "Maps/Directions 키가 없어도 지도·교통 추정은 graceful degrade 됩니다."
            ),
        },
        {"kind": "h3", "text": "2.3 모바일 환경 변수 (`apps/mobile/.env`)"},
        {
            "kind": "code",
            "text": (
                "# Maps SDK용 모바일 키 (권장: Android 패키지 com.nineru.trip + SHA-1 제한)\n"
                "EXPO_PUBLIC_GOOGLE_MAPS_API_KEY=your_mobile_maps_key\n"
                "\n"
                "# API 주소 오버라이드 (선택). 미설정 시 에뮬 기본값 사용\n"
                "# EXPO_PUBLIC_API_BASE_URL=http://10.0.2.2:3011\n"
            ),
        },
        {"kind": "h3", "text": "2.4 API Base URL"},
        {
            "kind": "table",
            "headers": ["환경", "URL"],
            "rows": [
                ["Android 에뮬레이터", "`http://10.0.2.2:3011` (앱 기본값)"],
                ["iOS 시뮬 / 웹", "`http://localhost:3011`"],
                ["실기기", "PC LAN IP 예: `http://192.168.x.x:3011`"],
            ],
        },
        {
            "kind": "p",
            "text": "앱 홈 → **설정(API)** 에서 주소 변경·헬스체크가 가능합니다.",
        },
        {"kind": "h3", "text": "2.5 헬스 확인"},
        {
            "kind": "ol",
            "items": [
                "브라우저/curl: `GET http://localhost:3011/health`",
                "`geminiConfigured` / `googleMapsConfigured` / `wordpressConfigured` 플래그 확인",
                "앱 설정 화면의 **헬스 체크** 버튼으로도 동일 확인 가능",
            ],
        },
        # ---- 3 ----
        {"kind": "h2", "text": "3. 첫 실행 — 온보딩 · 설정"},
        {"kind": "h3", "text": "3.1 온보딩 (3단계)"},
        {
            "kind": "p",
            "text": (
                "첫 실행 시 온보딩 모달이 표시됩니다(저장 키 `@9rutrip/onboardingSeen`). "
                "**다음**으로 넘기고 마지막에 **시작하기**를 누릅니다."
            ),
        },
        {
            "kind": "ol",
            "items": [
                "**여행 만들기** — 홈에서 「새 여행 만들기」로 도쿄·오사카 일정. 카드 ⋯ 로 삭제·복제.",
                "**일정 다듬기** — ≡ 길게 순서, 🕒 시각, Day▶ 다른 날 이동, 하단 「실행 취소」.",
                "**현장 · 한 손 조작** — 여행 시작 후 「다음으로 갈 곳」에서 길안내·완료·재루트.",
            ],
        },
        {"kind": "h3", "text": "3.2 설정 — 테마 · API"},
        {
            "kind": "ul",
            "items": [
                "**테마**: 시스템 / 라이트 / 다크 (Capture·Expenses·Summary·현장 UI까지 토큰 적용).",
                "**API 주소**: 에뮬은 `http://10.0.2.2:3011`, 실기기는 LAN IP.",
                "**헬스 체크** → 저장. 연결 실패 시 API 프로세스·방화벽·주소를 확인하세요.",
            ],
        },
        # ---- 4 ----
        {"kind": "h2", "text": "4. 여행 만들기"},
        {
            "kind": "ol",
            "items": [
                "홈에서 **새 여행 만들기**를 탭합니다.",
                "**도시**: 도쿄 또는 오사카 선택. (멀티시티는 Plan에서 Day별 도시 칩으로 배정 — 5장 B8 참고)",
                "**박수 / 일수 / 인원**을 입력합니다.",
                "**AI로 일정 생성**을 누릅니다. Gemini 실패 시 도쿄 폴백 코스가 사용됩니다.",
                "생성되면 Plan(일정) 화면으로 이동합니다. 숙소 후보(`lodgingCandidates`)가 함께 올 수 있습니다.",
            ],
        },
        # ---- 5 ----
        {"kind": "h2", "text": "5. Plan 일정 — 다듬기"},
        {
            "kind": "p",
            "text": (
                "Plan은 Day 탭 + 상단 압축 지도(~높이 37%) + DnD 리스트 구성입니다. "
                "마커와 리스트 선택이 연동됩니다. Maps 키가 없어도 힌트와 함께 동작합니다."
            ),
        },
        {"kind": "h3", "text": "5.1 Day · 필터 · 장소 추가"},
        {
            "kind": "ul",
            "items": [
                "상단 **Day1 / Day2 / …** 탭으로 당일을 고릅니다.",
                "카테고리 칩 **전체 | 맛집 | 관광 | 숙소** 로 목록을 좁힙니다.",
                "필터 ON 상태에서도 ≡ DnD로 필터된 항목만 재정렬하면, 당일 전체 시퀀스에 splice 됩니다(다른 카테고리 상대 위치 유지).",
                "**+추가** → 후보 모달. Places 응답이 있으면 카드에 **Google Places** 뱃지, 없으면 정적 POI.",
            ],
        },
        {"kind": "h3", "text": "5.2 ≡ DnD · 스와이프 삭제 · Undo"},
        {
            "kind": "ul",
            "items": [
                "**≡ 핸들만** 길게 눌러 드래그합니다(행 전체 드래그 아님). `activationDistance`로 오작동을 줄입니다.",
                "**왼쪽 스와이프**로 삭제(확인 후 enrich). 드래그 중에는 스와이프가 꺼집니다.",
                "행의 **삭제** 버튼으로도 동일하게 제거 가능합니다.",
                "DnD / Day 이동 / 삭제 직후 하단 **실행 취소** 스낵바(~5초). Undo 스택 깊이 N=5, 깊이 표시 가능.",
                "순서 변경 시 로컬에 즉시 반영되고 **교통 재계산 중…** 인디케이터가 뜹니다(낙관적 enrich).",
            ],
        },
        {"kind": "h3", "text": "5.3 Day▶ · 예정 시각"},
        {
            "kind": "ul",
            "items": [
                "행의 **Day▶** 로 다른 `dayIndex`로 옮긴 뒤 순서를 다시 매기고 enrich 합니다.",
                "**🕒** 탭 → HH:mm 예정 시각 편집.",
            ],
        },
        {"kind": "h3", "text": "5.4 지도 순서 모드 (B7)"},
        {
            "kind": "ul",
            "items": [
                "지도 마커를 **롱프레스**하면 **순서 모드**가 켜집니다.",
                "지도 하단 **≡ 번호 스트립**: 길게 누른 뒤 다른 번호를 탭하거나 **▲▼** 로 순서 변경.",
                "리스트·enrich·Undo와 동기화됩니다.",
                "참고: Android `react-native-maps` 핀 geo-드래그는 불안정하여 **비활성**입니다. 순서 모드/스트립이 대체 UX입니다.",
            ],
        },
        {"kind": "h3", "text": "5.5 이동 · 비교 · Yahoo/Google 환승"},
        {
            "kind": "ol",
            "items": [
                "장소 행의 **이동 · 비교 ›** 칩을 탭합니다.",
                "도보 / 대중교통 / 택시의 분·비용을 비교하는 시트가 열립니다.",
                "선호 수단(`preferredTransportMode`)을 고르면 glance·`travelFromPrev*`에 반영됩니다.",
                "Maps 서버 키가 있으면 Directions(walking/driving) 실측, 없으면 모드별 하버사인.",
                "**일본 대중교통**: Google Directions API는 JP transit 미지원 → `haversine:transit`(또는 파트너)이 정상입니다.",
                "시트에서 **Yahoo 환승** / **Google 환승** 딥링크로 외부 앱에서 정확한 환승을 확인하세요.",
                "`JP_TRANSIT_PROVIDER=navitime` + 자격증명이 있으면 `partner:navitime` 시도를 합니다(실패 시 추정+딥링크 폴백).",
            ],
        },
        {"kind": "h3", "text": "5.6 숙소"},
        {
            "kind": "ul",
            "items": [
                "여행 설정(⋯ 더보기) 또는 숙소 후보에서 Top N을 고릅니다.",
                "점수 설명: **역세권 근접 / 가격 경쟁력 / 평점** (scoreBreakdown 한국어 라인).",
                "숙소 선택 시 동선(교통)이 다시 계산됩니다.",
            ],
        },
        {"kind": "h3", "text": "5.7 동선 최적화"},
        {
            "kind": "ol",
            "items": [
                "당일 장소가 충분할 때 **동선 최적화**를 누릅니다.",
                "`POST /trip/optimize-day` — Gemini 재배치, 실패 시 nearest-neighbor 폴백.",
                "미리보기 Alert에서 적용/취소. 적용 후 enrich.",
            ],
        },
        {"kind": "h3", "text": "5.8 멀티시티 Day 배정 (B8)"},
        {
            "kind": "p",
            "text": (
                "여행에 여러 도시가 있을 때 Day별 **도쿄/오사카** 칩으로 `cities.dayIndexes`를 수동 할당합니다. "
                "지도 중심이 갱신되고, (선택) 당일 `place.cityId` 갱신 프롬프트가 뜰 수 있습니다."
            ),
        },
        {"kind": "h3", "text": "5.9 체크리스트 · 날씨"},
        {
            "kind": "ul",
            "items": [
                "**체크인 체크리스트**: 예약번호·여권·WiFi·미팅포인트 등 — Trip에 저장. Plan ⋯ / 현장 모드에서 체크박스.",
                "**오늘 날씨**: Open-Meteo(키 없음) → 「오늘 날씨 · °C · 강수%」 칩 + 시간대 혼잡 휴리스틱.",
            ],
        },
        {"kind": "h3", "text": "5.10 여행 설정 · 하단 이동"},
        {
            "kind": "ul",
            "items": [
                "**⋯ 더보기 · 여행 설정**: 가이드알람 ON/OFF, AI재루트 ON/OFF, 숙소 후보 등.",
                "하단: 지도 / 캡처 / 경비 / 요약 등 관련 화면으로 이동.",
                "**여행 시작**으로 상태를 `active`(여행중)로 바꿉니다.",
            ],
        },
        # ---- 6 ----
        {"kind": "h2", "text": "6. 여행 중 — 현장 모드"},
        {
            "kind": "p",
            "text": (
                "`status === active` 이면 **현장** / **일정** 탭을 쓸 수 있습니다. "
                "현장 모드에서는 「다음으로 갈 곳」 패널이 크게 표시됩니다."
            ),
        },
        {"kind": "h3", "text": "6.1 길안내 · 완료 · 재루트"},
        {
            "kind": "ul",
            "items": [
                "**길안내** — Google Maps 내비게이션(Android `google.navigation` 또는 https 경로, transit 가능).",
                "**완료** — 해당 장소를 완료 처리(`completedPlaceIds`). 다음 미완료 장소로 배너가 갱신됩니다.",
                "**일정 재루트** — AI재루트 ON일 때 `POST /trip/reroute`. 응답 후 **재루트 미리보기** Alert(변경 요약) → 적용/취소.",
                "AI재루트가 OFF면 「아래에서 AI 재루트를 켠 뒤 다시 시도」 안내가 뜹니다.",
            ],
        },
        {"kind": "h3", "text": "6.2 가이드 알람"},
        {
            "kind": "p",
            "text": (
                "가이드알람 ON + 여행중이면 다음 미완료 장소 배너와, 임박 시 인앱 Alert · "
                "Android 로컬 알림이 동작합니다."
            ),
        },
        {"kind": "h3", "text": "6.3 GPS 이탈"},
        {
            "kind": "p",
            "text": (
                "여행중 + AI재루트 ON일 때 GPS가 경로에서 벗어나면 "
                "「경로에서 벗어난 것 같아요 — 재루트할까요?」 배너가 뜹니다. **재루트**로 미리보기 흐름에 진입합니다."
            ),
        },
        # ---- 7 ----
        {"kind": "h2", "text": "7. 사진·리뷰 / 경비 / 요약·WordPress"},
        {"kind": "h3", "text": "7.1 사진 · 리뷰 (Capture)"},
        {
            "kind": "ul",
            "items": [
                "여행 중(또는 Plan에서) **캡처**로 장소별 사진+리뷰를 남깁니다.",
                "데이터는 9ruDocs 호환 **Step** 스키마로 저장되어 이후 발행에 쓰입니다.",
            ],
        },
        {"kind": "h3", "text": "7.2 경비 (Expenses)"},
        {
            "kind": "ul",
            "items": [
                "**현금/수동**: 항목명 + 금액(엔) + 카테고리 입력 후 추가.",
                "**SMS 붙여넣기**: 카드 결제 SMS를 복사 → **클립보드 붙여넣기** → **SMS 파싱** "
                "(`POST /trip/parse-sms` + 로컬 폴백). 가맹점·KRW·추정 JPY 확인 후 추가.",
                "인박스 자동 읽기는 커스텀 빌드가 필요하며 기본 Expo Go 경로에는 없습니다(`apps/mobile/docs/SMS.md`).",
            ],
        },
        {"kind": "h3", "text": "7.3 요약 · 인사이트 · WordPress (Summary)"},
        {
            "kind": "ul",
            "items": [
                "계획 vs 실제 비용 요약, 식비/교통 비중 등 **한국어 인사이트** 1–3줄.",
                "**발행용 초안 내보내기** — `POST /trip/export-draft` → BlogDraft 호환 JSON (리뷰 1개 이상 권장).",
                "**WordPress 발행** — `POST /wordpress/publish` (Application Password). 사이트 URL·계정은 API `.env`.",
                "리뷰가 없으면 내보내기/발행 전 「먼저 사진/리뷰를 캡처」 안내가 뜹니다.",
            ],
        },
        # ---- 8 ----
        {"kind": "h2", "text": "8. 홈 — 상태 칩 · 삭제 · 복제"},
        {
            "kind": "ul",
            "items": [
                "저장된 여행 카드에 **계획중 / 여행중 / 완료** 상태 칩이 보입니다.",
                "카드를 **길게 누르기** 또는 **⋯** 메뉴 → **삭제** / **복제**.",
                "삭제: 확인 후 되돌릴 수 없음.",
                "복제: 새 id, 상태는 계획중, 완료/리뷰/경비는 초기화(일정·숙소 후보는 유지).",
                "브랜드 히어로 아래 **새 여행 만들기** · **API 설정** 진입.",
            ],
        },
        # ---- 9 ----
        {"kind": "h2", "text": "9. FAQ"},
        {"kind": "h3", "text": "Q. Google Maps 키가 없어도 되나요?"},
        {
            "kind": "p",
            "text": (
                "됩니다. 지도는 힌트/제한 모드로 열리고, 교통은 모드별 하버사인 폴백입니다. "
                "실측 Directions·지도 타일이 필요하면 `docs/GOOGLE-MAPS-API-KEY.md`를 따라 "
                "모바일 키(`EXPO_PUBLIC_GOOGLE_MAPS_API_KEY`)와 서버 키(`GOOGLE_MAPS_API_KEY`)를 넣으세요."
            ),
        },
        {"kind": "h3", "text": "Q. 일본 대중교통(환승) 시간이 왜 ‘추정’인가요?"},
        {
            "kind": "p",
            "text": (
                "Google Maps Platform Directions/Routes는 일본 transit 파트너를 지원하지 않습니다. "
                "키를 넣어도 transit은 `haversine:transit`이 정상인 경우가 많습니다. "
                "정확한 환승은 비교 시트의 **Yahoo 환승 / Google 환승** 딥링크를 쓰세요. "
                "NAVITIME 계약 시 `partner:navitime`을 시도합니다 — `docs/JP-TRANSIT-PARTNER.md`."
            ),
        },
        {"kind": "h3", "text": "Q. Places 뱃지가 안 보여요"},
        {
            "kind": "p",
            "text": (
                "장소 추가는 Places Text가 가능할 때 `source=places`와 **Google Places** 뱃지를 표시합니다. "
                "키가 없거나 Places API가 비활성이면 정적 POI 후보만 나옵니다."
            ),
        },
        {"kind": "h3", "text": "Q. 에뮬에서 API 연결이 안 됩니다"},
        {
            "kind": "ol",
            "items": [
                "PC에서 `npm run api`가 떠 있는지, `GET http://localhost:3011/health` 확인.",
                "앱 API 주소가 **`http://10.0.2.2:3011`** 인지 확인(`localhost`는 에뮬 자기 자신).",
                "방화벽·포트 3011 허용, Expo 재시작.",
            ],
        },
        {"kind": "h3", "text": "Q. ≡ 드래그가 안 되거나 스크롤과 싸웁니다"},
        {
            "kind": "p",
            "text": (
                "행이 아니라 **≡ 핸들만** 길게 누르세요. 왼쪽 스와이프는 삭제용이며, "
                "드래그 중에는 스와이프가 꺼집니다. 시스템 접근성 **동작 줄이기(reduce motion)** 가 켜져 있으면 "
                "애니메이션이 약해질 수 있습니다."
            ),
        },
        {"kind": "h3", "text": "Q. WordPress 발행이 실패합니다"},
        {
            "kind": "p",
            "text": (
                "`WP_SITE_URL`(또는 `WP_BASE_URL`), `WP_USERNAME`, `WP_APP_PASSWORD`를 API `.env`에 설정하고 "
                "API를 재시작하세요. `/health`의 `wordpressConfigured`가 true여야 합니다. "
                "리뷰(캡처)가 최소 1개 있어야 합니다."
            ),
        },
        # ---- 10 ----
        {"kind": "h2", "text": "10. 부록 — 참고 문서"},
        {
            "kind": "ul",
            "items": [
                "[`docs/GOOGLE-MAPS-API-KEY.md`](./GOOGLE-MAPS-API-KEY.md) — Maps 키 발급·제한·검증",
                "[`docs/JP-TRANSIT-PARTNER.md`](./JP-TRANSIT-PARTNER.md) — NAVITIME / Yahoo·Google 환승 딥링크",
                "[`docs/E2E-CHECKLIST.md`](./E2E-CHECKLIST.md) — Android 실기기 전체 사이클",
                "[`docs/9ruTrip-제품개요.md`](./9ruTrip-제품개요.md) — 제품 개요(슬라이드 동반)",
                "[`apps/mobile/docs/SMS.md`](../apps/mobile/docs/SMS.md) — SMS 경비·공유 인텐트 메모",
                "저장소 README — 아키텍처·API 표·P0–P3 / Sprint A·B 요약",
            ],
        },
        {
            "kind": "p",
            "text": (
                "문서 재생성: "
                "`C:\\Users\\Intellian\\AppData\\Local\\Programs\\Python\\Python312\\python.exe "
                "docs\\generate_user_manual_docx.py`"
            ),
        },
    ]


# ---------------------------------------------------------------------------
# Markdown writer
# ---------------------------------------------------------------------------


def _md_escape_cell(s: str) -> str:
    return s.replace("|", "\\|").replace("\n", " ")


def blocks_to_markdown(blocks: list[ContentBlock]) -> str:
    lines: list[str] = []
    for b in blocks:
        kind = b["kind"]
        if kind == "h1":
            lines.append(f"# {b['text']}")
            lines.append("")
        elif kind == "h2":
            lines.append(f"## {b['text']}")
            lines.append("")
        elif kind == "h3":
            lines.append(f"### {b['text']}")
            lines.append("")
        elif kind == "p":
            lines.append(b["text"])
            lines.append("")
        elif kind == "note":
            lines.append(f"> **참고:** {b['text']}")
            lines.append("")
        elif kind == "ul":
            for item in b["items"]:
                lines.append(f"- {item}")
            lines.append("")
        elif kind == "ol":
            for i, item in enumerate(b["items"], 1):
                lines.append(f"{i}. {item}")
            lines.append("")
        elif kind == "code":
            lines.append("```")
            lines.append(b["text"].rstrip("\n"))
            lines.append("```")
            lines.append("")
        elif kind == "table":
            headers = b["headers"]
            rows = b["rows"]
            lines.append("| " + " | ".join(_md_escape_cell(h) for h in headers) + " |")
            lines.append("| " + " | ".join("---" for _ in headers) + " |")
            for row in rows:
                lines.append(
                    "| " + " | ".join(_md_escape_cell(c) for c in row) + " |"
                )
            lines.append("")
        else:
            raise ValueError(f"unknown block kind: {kind}")
    return "\n".join(lines).rstrip() + "\n"


# ---------------------------------------------------------------------------
# DOCX writer
# ---------------------------------------------------------------------------


def _set_run_font(run, *, size: int, bold: bool = False, color: RGBColor = C_INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    # East Asian font (Malgun Gothic on Windows Word)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)


def _add_runs_with_bold(paragraph, text: str, *, size: int, color: RGBColor = C_INK):
    """Simple **bold** markers inside a paragraph."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        _set_run_font(run, size=size, bold=(i % 2 == 1), color=color)


def _style_paragraph(p, *, space_after: int = 8):
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)


def write_docx(blocks: list[ContentBlock], path: Path) -> None:
    doc = Document()
    # Default style font
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for b in blocks:
        kind = b["kind"]
        if kind == "h1":
            p = doc.add_heading(b["text"], level=0)
            for run in p.runs:
                _set_run_font(run, size=22, bold=True, color=C_PRIMARY)
            _style_paragraph(p, space_after=12)
        elif kind == "h2":
            p = doc.add_heading(b["text"], level=1)
            for run in p.runs:
                _set_run_font(run, size=16, bold=True, color=C_PRIMARY)
            _style_paragraph(p, space_after=10)
        elif kind == "h3":
            p = doc.add_heading(b["text"], level=2)
            for run in p.runs:
                _set_run_font(run, size=13, bold=True, color=C_PRIMARY)
            _style_paragraph(p, space_after=6)
        elif kind == "p":
            for line in b["text"].split("\n"):
                p = doc.add_paragraph()
                _add_runs_with_bold(p, line, size=11)
                _style_paragraph(p, space_after=6)
        elif kind == "note":
            p = doc.add_paragraph()
            run = p.add_run("참고: ")
            _set_run_font(run, size=10, bold=True, color=C_MUTED)
            _add_runs_with_bold(p, b["text"], size=10, color=C_MUTED)
            _style_paragraph(p, space_after=10)
        elif kind in ("ul", "ol"):
            style_name = "List Bullet" if kind == "ul" else "List Number"
            for item in b["items"]:
                p = doc.add_paragraph(style=style_name)
                # clear default empty run if any
                if p.runs:
                    p.runs[0].text = ""
                _add_runs_with_bold(p, item, size=11)
                _style_paragraph(p, space_after=4)
        elif kind == "code":
            for line in b["text"].rstrip("\n").split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(line if line else " ")
                _set_run_font(run, size=9, color=C_INK)
                run.font.name = "Consolas"
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    rFonts.set(qn("w:ascii"), "Consolas")
                    rFonts.set(qn("w:hAnsi"), "Consolas")
                    rFonts.set(qn("w:eastAsia"), FONT)
                pf = p.paragraph_format
                pf.space_after = Pt(0)
                pf.space_before = Pt(0)
                pf.left_indent = Pt(12)
            # spacer after code block
            sp = doc.add_paragraph()
            _style_paragraph(sp, space_after=8)
        elif kind == "table":
            headers = b["headers"]
            rows = b["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                _set_run_font(run, size=10, bold=True, color=C_PRIMARY)
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    # strip markdown backticks for word readability
                    plain = cell_text.replace("`", "")
                    run = p.add_run(plain)
                    _set_run_font(run, size=10)
            doc.add_paragraph()
        else:
            raise ValueError(f"unknown block kind: {kind}")

    # Footer note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"9ruTrip · 사용자 매뉴얼 · brand #0c4a6e · ref {REF_COMMIT}"
    )
    _set_run_font(run, size=8, color=C_MUTED)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    blocks = build_blocks()
    md = blocks_to_markdown(blocks)
    OUT_MD.write_text(md, encoding="utf-8")
    write_docx(blocks, OUT_DOCX)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    print(f"MD bytes: {OUT_MD.stat().st_size}")
    print(f"DOCX bytes: {OUT_DOCX.stat().st_size}")


if __name__ == "__main__":
    main()
